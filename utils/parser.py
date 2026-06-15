from io import BytesIO
from pathlib import Path

import pdfplumber


def _get_file_bytes(uploaded_file) -> bytes:
    if uploaded_file is None:
        return b""

    if hasattr(uploaded_file, "getvalue"):
        return uploaded_file.getvalue()

    if hasattr(uploaded_file, "read"):
        current_position = None
        if hasattr(uploaded_file, "tell"):
            current_position = uploaded_file.tell()
        if hasattr(uploaded_file, "seek"):
            uploaded_file.seek(0)
        data = uploaded_file.read()
        if current_position is not None and hasattr(uploaded_file, "seek"):
            uploaded_file.seek(current_position)
        return data

    raise ValueError("Unsupported uploaded file type.")


def extract_text_from_pdf(uploaded_file):
    if uploaded_file is None:
        return ""

    file_bytes = _get_file_bytes(uploaded_file)
    if not file_bytes:
        return ""

    text = ""
    try:
        with pdfplumber.open(BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception:
        return ""

    return text


def extract_text_from_docx(uploaded_file):
    if uploaded_file is None:
        return ""

    file_bytes = _get_file_bytes(uploaded_file)
    try:
        from docx import Document as DocxDocument
    except ImportError as error:
        raise RuntimeError("Install python-docx to parse DOCX resumes.") from error

    doc = DocxDocument(BytesIO(file_bytes))
    paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
    table_cells = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    table_cells.append(cell.text)

    return "\n".join(paragraphs + table_cells)


def extract_resume_text(uploaded_file, filename=None):
    suffix = Path(filename or getattr(uploaded_file, "name", "")).suffix.lower()
    if suffix == ".docx":
        return extract_text_from_docx(uploaded_file)
    if suffix == ".pdf":
        return extract_text_from_pdf(uploaded_file)
    raise ValueError("Unsupported resume format. Upload a PDF or DOCX file.")
